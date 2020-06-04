from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from datetime import datetime
current_time = datetime.now()
date = current_time.strftime('%m/%d/%Y')

company = input("Company name? ")
manager = input("is there a manager mentioned? ")

#adds the heading info
document = Document()

style = document.styles['Normal']
font = style.font
font.name = 'Cambria Math'
font.size = Pt(12)

document.add_paragraph('Gabriel Ng')
document.add_paragraph('(551)-587-1056')
document.add_paragraph('gabeng15@gmail.com')
document.add_paragraph(date)
#if there is a manager mentioned he will be in the greetings
if manager == "" :
    document.add_paragraph('Dear Hiring Manager,')
else :
    document.add_paragraph('Dear ' + manager)
#first paragraph
p = document.add_paragraph('I am excited to be applying as a software engineer at ' + company + '.')
p.add_run(' As a newly graduate from Rutgers University, I will use what I learned from my time there to help build the quality products you need.')
p.add_run(' I am detailed orientated, and always thinking of ways to improve my code and myself.')
#body paragraphs
document.add_paragraph('During my studies in Rutgers University, I\'ve learned many things about Computer Science, learning, and about other and different people. I learned the most about Computer Science and programming through hands on projects, like how my Python interpreter in C++ taught me about how programming languages work. ')
document.add_paragraph('Another great project is my group project that put me and 2 other students, to create an automatic exam grader for python functions. In this project, I acted as a backend engineer developing a database in MySQL, using phpMyAdmin, while my other 2 members acted as the front end and middle end. Since we needed to meet up frequently, I took up the mantle of group leader and made sure everyone was contributing their part. In the end we all finished an A in the class.')
#end paragraph
p = document.add_paragraph('Thank you for your time and consideration. I\'m looking forward to learning more details about the Software Engineering position at ' + company + '.')
p.add_run(' I\'m excited about the opportunity to apply my skills and develop them.')
document.add_paragraph('Sincerely,')
document.add_paragraph('Gabriel Ng')

document.save('Gabriel Ng_Cover Letter.docx')